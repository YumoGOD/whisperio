from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document


def format_timestamp_srt(seconds: float) -> str:
    milliseconds = int(round(seconds * 1000))
    hours, remainder = divmod(milliseconds, 3_600_000)
    minutes, remainder = divmod(remainder, 60_000)
    secs, millis = divmod(remainder, 1000)
    return f"{hours:02}:{minutes:02}:{secs:02},{millis:03}"


def format_timestamp_vtt(seconds: float) -> str:
    return format_timestamp_srt(seconds).replace(",", ".")


def format_timestamp_docx(seconds: float | None) -> str:
    total_seconds = int(round(max(0.0, float(seconds or 0))))
    hours, remainder = divmod(total_seconds, 3600)
    minutes, secs = divmod(remainder, 60)
    return f"{hours:02}:{minutes:02}:{secs:02}"


def segments_to_srt(segments: list[dict[str, Any]]) -> str:
    blocks = []
    for index, segment in enumerate(segments, start=1):
        blocks.append(
            "\n".join(
                [
                    str(index),
                    f"{format_timestamp_srt(segment['start'])} --> {format_timestamp_srt(segment['end'])}",
                    segment.get("text", "").strip(),
                ]
            )
        )
    return "\n\n".join(blocks) + ("\n" if blocks else "")


def segments_to_vtt(segments: list[dict[str, Any]]) -> str:
    body = []
    for segment in segments:
        body.append(
            "\n".join(
                [
                    f"{format_timestamp_vtt(segment['start'])} --> {format_timestamp_vtt(segment['end'])}",
                    segment.get("text", "").strip(),
                ]
            )
        )
    return "WEBVTT\n\n" + "\n\n".join(body) + ("\n" if body else "")


def write_docx_export(path: Path, segments: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for index, segment in enumerate(segments):
        start = format_timestamp_docx(segment.get("start"))
        end = format_timestamp_docx(segment.get("end"))
        timing = document.add_paragraph()
        timing.add_run(f"({start}-{end})").bold = True
        document.add_paragraph(str(segment.get("text") or "").strip())
        if index != len(segments) - 1:
            document.add_paragraph()
    document.save(path)


def write_exports(
    transcript_dir: Path,
    job_id: str,
    text: str,
    segments: list[dict[str, Any]],
) -> dict[str, Path]:
    transcript_dir.mkdir(parents=True, exist_ok=True)
    paths = {
        "txt": transcript_dir / f"{job_id}.txt",
        "json": transcript_dir / f"{job_id}.json",
        "srt": transcript_dir / f"{job_id}.srt",
        "vtt": transcript_dir / f"{job_id}.vtt",
        "docx": transcript_dir / f"{job_id}.docx",
    }
    paths["txt"].write_text(text.strip() + "\n", encoding="utf-8")
    paths["srt"].write_text(segments_to_srt(segments), encoding="utf-8")
    paths["vtt"].write_text(segments_to_vtt(segments), encoding="utf-8")
    write_docx_export(paths["docx"], segments)
    # JSON is intentionally NOT written here — pipeline.py writes it once
    # after all stage timings (export_seconds, elapsed_seconds, rtf) are finalized.
    return paths
